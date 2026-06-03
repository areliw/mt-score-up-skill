#!/usr/bin/env python3
"""
One-time programmatic marking of ตัวอย่าง B template with docxtpl placeholders.

Reads the converted .docx template (output of .doc → .docx conversion via Word COM),
locates known cell positions in the section header table, and replaces literal text
with Jinja-style placeholders (`{{ doc_subject }}`, `{{ doc_code }}`, etc.).

For OTHER hospitals' templates, do this manually in Word instead — open the .docx,
find each cell holding subject/code/page/revision data, and replace the literal
value with the corresponding `{{ placeholder }}`. The exact cell location doesn't
matter — docxtpl finds placeholders wherever they are.

Usage:
    python scripts/mark_template.py <input.docx> <output_marked.docx>
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass


def mark_hospital_b_template(src: Path, dst: Path) -> None:
    """Programmatically add docxtpl placeholders at known ตัวอย่าง B positions.

    Section header structure (3 rows × 1 col):
      Row 0: hospital banner       → no change
      Row 1: 3 paragraphs            → mark subject + doc_code + page
      Row 2: 1 paragraph             → mark rev_no + rev_date
    """
    doc = Document(src)
    if not doc.sections or not doc.sections[0].header.tables:
        raise RuntimeError("Template has no section header tables — wrong format")

    tbl = doc.sections[0].header.tables[0]

    # Row 1: subject (p[0]), continuation (p[1] cleared), doc_code+page (p[2])
    if len(tbl.rows) >= 2:
        cell = tbl.rows[1].cells[0]
        paras = cell.paragraphs

        if len(paras) >= 1:
            paras[0].text = "วิธีปฏิบัติเรื่อง : {{ doc_subject }}"
        if len(paras) >= 2:
            paras[1].text = ""  # subject continuation — cleared (no overflow needed)
        if len(paras) >= 3:
            paras[2].text = (
                "รหัสเอกสาร   : {{ doc_code }}"
                "                                                            "
                "หน้าที่ :{{ page }}/{{ total_pages }}"
            )

    # Row 2: revision
    if len(tbl.rows) >= 3:
        cell = tbl.rows[2].cells[0]
        if cell.paragraphs:
            cell.paragraphs[0].text = (
                "ทบทวนแก้ไขครั้งที่ : {{ rev_no }} วันที่ {{ rev_date }}"
            )

    doc.save(dst)
    print(f"Marked template saved: {dst}")
    print("Placeholders added: doc_subject, doc_code, page, total_pages, rev_no, rev_date")


def main() -> int:
    if len(sys.argv) != 3:
        print(__doc__)
        return 2
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    if not src.exists():
        print(f"ERROR: {src} not found")
        return 1
    mark_hospital_b_template(src, dst)
    return 0


if __name__ == "__main__":
    sys.exit(main())
