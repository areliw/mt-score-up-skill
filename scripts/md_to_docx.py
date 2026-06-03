#!/usr/bin/env python3
"""
Minimal Markdown -> DOCX converter for WI files (Thai-aware).

Supports YAML frontmatter at top of markdown for hospital header + doc control box:

    ---
    hospital_org: กลุ่มงานเทคนิคการแพทย์และพยาธิวิทยาคลินิก โรงพยาบาลตัวอย่าง B
    doc_subject: <full title of WI>
    doc_code: WI-BBCT-040
    page: 1
    total_pages: 7
    revision_no: 0
    revision_date: 7 มกราคม 2569
    ---
    # Body markdown...

If frontmatter present, builds:
- Hospital banner table (logo + org name)
- Document control table (subject / code+page / revision)
At top of doc. Then renders markdown body as before.

Body parsing:
- `## heading` -> Heading 1, `### heading` -> Heading 2, `#### heading` -> Heading 3
- `| a | b | c |` table blocks -> Word table (Table Grid)
- `- item` -> bullet list
- `**text**` inline bold (preserved)
- Other lines -> plain paragraph

Usage:
    python scripts/md_to_docx.py <input.md> <output.docx>
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


THAI_FONT = "TH Sarabun New"


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = THAI_FONT
    style.font.size = Pt(14)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(4)  # separation via paragraph spacing, not blank paras
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), THAI_FONT)
    rfonts.set(qn("w:hAnsi"), THAI_FONT)
    rfonts.set(qn("w:cs"), THAI_FONT)

    # Heading styles → BLACK bold TH Sarabun, compact spacing (Word's default is blue +
    # large + airy; Thai gov/hospital docs use black bold headings tight to the body).
    for level, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)):
        hs = doc.styles[level]
        hs.font.name = THAI_FONT
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hrpr = hs.element.get_or_add_rPr()
        hrf = hrpr.find(qn("w:rFonts"))
        if hrf is None:
            hrf = OxmlElement("w:rFonts")
            hrpr.append(hrf)
        hrf.set(qn("w:ascii"), THAI_FONT)
        hrf.set(qn("w:hAnsi"), THAI_FONT)
        hrf.set(qn("w:cs"), THAI_FONT)
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(3)
        hs.paragraph_format.keep_with_next = True


def _unquote(s: str) -> str:
    s = s.strip()
    if len(s) >= 2 and s[0] == s[-1] and s[0] in ("'", '"'):
        return s[1:-1]
    return s


def parse_simple_yaml(fm_text: str) -> dict:
    """Parse flat YAML: scalar `key: value` + single-level string lists (`key:` then `  - item`)."""
    meta: dict = {}
    current_list_key = None
    for line in fm_text.splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        if stripped.startswith("- ") and current_list_key:
            meta[current_list_key].append(_unquote(stripped[2:]))
            continue
        if ":" in line:
            key, _, value = line.partition(":")
            key, value = key.strip(), value.strip()
            if value == "":
                meta[key] = []
                current_list_key = key
            else:
                meta[key] = _unquote(value)
                current_list_key = None
    return meta


def parse_frontmatter(text: str) -> tuple[dict, str]:
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 4)
    if end == -1:
        return {}, text
    return parse_simple_yaml(text[4:end].strip()), text[end + 4 :].lstrip("\n")


def load_profile(slug: str) -> dict | None:
    """Load a per-hospital layout profile from profiles/<slug>.yaml. Returns None if absent."""
    p = Path(__file__).parent.parent / "profiles" / f"{slug}.yaml"
    if not p.exists():
        return None
    return parse_simple_yaml(p.read_text(encoding="utf-8"))


def _subst(template: str, values: dict) -> str:
    """Replace {field} in template with values[field] (or [field] placeholder if missing)."""
    return re.sub(
        r"\{(\w+)\}",
        lambda m: str(values.get(m.group(1), f"[{m.group(1)}]")),
        template,
    )


def apply_hospital_profile(meta: dict) -> None:
    """If meta has `hospital: <slug>`, load that profile and resolve its header template
    using canonical fields → sets meta[header_logo/header_org/header_rows].

    This is the SCALE mechanism: WI frontmatter = canonical content (hospital-agnostic);
    profile = per-hospital layout (reused). Same WI renders in any hospital's layout by
    swapping `hospital:`.
    """
    slug = meta.get("hospital")
    if not slug:
        return
    prof = load_profile(slug)
    if not prof:
        return
    values = {**prof, **meta}  # canonical (meta) overrides profile defaults
    if prof.get("logo"):
        meta["header_logo"] = prof["logo"]
    if prof.get("org"):
        meta["header_org"] = prof["org"]
    if prof.get("header"):
        meta["header_rows"] = [_subst(line, values) for line in prof["header"]]


def _span_row(table, row_idx: int) -> None:
    """Span all cells in row_idx into the first cell via OOXML gridSpan.

    This is the safe alternative to python-docx's cell.merge(), which can
    produce XML that Word treats as needing repair (the "Working on it..."
    hang). We set gridSpan on the first cell and physically remove the
    other cells from the row.
    """
    row = table.rows[row_idx]
    first_tc = row.cells[0]._tc
    extras = [c._tc for c in row.cells[1:]]
    num_cols = len(row.cells)

    tcPr = first_tc.get_or_add_tcPr()
    for gs in tcPr.findall(qn("w:gridSpan")):
        tcPr.remove(gs)
    grid_span = OxmlElement("w:gridSpan")
    grid_span.set(qn("w:val"), str(num_cols))
    tcPr.append(grid_span)

    for tc in extras:
        tc.getparent().remove(tc)


def build_header_tables(doc: Document, meta: dict[str, str]) -> None:
    """Build the hospital banner + doc control table at the top of the doc."""
    # 1. Hospital banner — 1 row, 2 col: [LOGO] | org_name
    if meta.get("hospital_org"):
        banner = doc.add_table(rows=1, cols=2)
        banner.style = "Table Grid"
        banner.autofit = False

        logo_cell = banner.cell(0, 0)
        logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        logo_p = logo_cell.paragraphs[0]
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Embed real logo image if logo_path given + file exists; else [LOGO] placeholder
        logo_file = None
        if meta.get("logo_path"):
            cand = Path(meta["logo_path"])
            if not cand.is_absolute():
                cand = Path(__file__).parent.parent / meta["logo_path"]
            if cand.exists():
                logo_file = str(cand)
        if logo_file:
            try:
                logo_p.add_run().add_picture(logo_file, width=Cm(2.2))
            except Exception:
                logo_p.add_run("[LOGO]").italic = True
        else:
            logo_p.add_run("[LOGO]").italic = True

        org_cell = banner.cell(0, 1)
        org_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        org_p = org_cell.paragraphs[0]
        org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org_run = org_p.add_run(meta["hospital_org"])
        org_run.bold = True
        org_run.font.size = Pt(16)

    # 2. Doc control box — 3 rows, 2 cols (rows 0 & 2 will gridSpan to look merged)
    if meta.get("doc_subject") or meta.get("doc_code"):
        ctrl = doc.add_table(rows=3, cols=2)
        ctrl.style = "Table Grid"

        # Row 1 (will span): วิธีปฏิบัติเรื่อง
        r1_p = ctrl.cell(0, 0).paragraphs[0]
        r1_p.add_run("วิธีปฏิบัติเรื่อง : ").bold = True
        r1_p.add_run(meta.get("doc_subject", "[TBD]"))

        # Row 2: รหัสเอกสาร | หน้าที่
        r2_left_p = ctrl.cell(1, 0).paragraphs[0]
        r2_left_p.add_run("รหัสเอกสาร : ").bold = True
        r2_left_p.add_run(meta.get("doc_code", "[TBD]"))

        r2_right_p = ctrl.cell(1, 1).paragraphs[0]
        r2_right_p.add_run("หน้าที่ : ").bold = True
        r2_right_p.add_run(
            f"{meta.get('page', '1')}/{meta.get('total_pages', '[X]')}"
        )

        # Row 3 (will span): ทบทวนแก้ไขครั้งที่
        r3_p = ctrl.cell(2, 0).paragraphs[0]
        r3_p.add_run("ทบทวนแก้ไขครั้งที่ : ").bold = True
        r3_p.add_run(
            f"{meta.get('revision_no', '0')} วันที่ {meta.get('revision_date', '[DD/MM/YY]')}"
        )

        # Apply spans AFTER content is set
        _span_row(ctrl, 0)
        _span_row(ctrl, 2)

    doc.add_paragraph("")  # spacing after header tables


_TRPR_ORDER = [
    "cnfStyle", "divId", "gridBefore", "gridAfter", "wBefore", "wAfter",
    "cantSplit", "trHeight", "tblHeader", "tblCellSpacing", "jc", "hidden",
]


def _keep_table_together(table) -> None:
    """Stop a table from breaking across a page boundary: forbid mid-row splits (cantSplit)
    and keep every row with the next, so Word moves the whole table to the next page if it
    doesn't fit instead of leaving a couple of orphan rows behind."""
    rows = table.rows
    for ri, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            _append_ordered(trPr, OxmlElement("w:cantSplit"), _TRPR_ORDER)
        if ri < len(rows) - 1:  # all but the last row "keep with next" → rows stay together
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True


def flush_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            table.cell(r_idx, c_idx).text = cell_text
    _keep_table_together(table)  # don't split this table across a page boundary


def is_table_separator(cells: list[str]) -> bool:
    return bool([c for c in cells if c.strip()]) and all(re.match(r"^:?-+:?$", c.strip()) for c in cells if c.strip())


def cells_from(line: str) -> list[str]:
    parts = [p.strip() for p in line.split("|")]
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return parts


def _add_rich_paragraph(doc, text: str, style: str | None = None):
    """Add a paragraph, rendering inline **bold** spans as REAL bold runs (not stripped)."""
    p = doc.add_paragraph(style=style)
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        if len(part) >= 4 and part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        else:
            p.add_run(part)
    return p


def render_body(doc: Document, body: str) -> None:
    table_buffer: list[list[str]] = []

    for raw in body.splitlines():
        line = raw.rstrip("\n")

        if line.lstrip().startswith("|") and "|" in line[1:]:
            cells = cells_from(line)
            if is_table_separator(cells):
                continue
            table_buffer.append(cells)
            continue

        if table_buffer:
            flush_table(doc, table_buffer)
            table_buffer = []

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("- ") or line.startswith("* "):
            _add_rich_paragraph(doc, line[2:].strip(), style="List Bullet")
        elif line.startswith("---"):
            continue  # markdown HR / frontmatter rule — no empty spacer paragraph
        else:
            if line.strip() == "":
                continue  # skip blank lines; Normal space_after handles separation
            _add_rich_paragraph(doc, line)

    if table_buffer:
        flush_table(doc, table_buffer)


def _merge_range(table, row_idx: int, start_col: int, end_col: int) -> None:
    """Merge cells [start_col..end_col] in a row via gridSpan (safe — no cell.merge hang)."""
    cells = table.rows[row_idx].cells
    first_tc = cells[start_col]._tc
    extras = [cells[c]._tc for c in range(start_col + 1, end_col + 1)]
    span = end_col - start_col + 1
    tcPr = first_tc.get_or_add_tcPr()
    for gs in tcPr.findall(qn("w:gridSpan")):
        tcPr.remove(gs)
    g = OxmlElement("w:gridSpan")
    g.set(qn("w:val"), str(span))
    tcPr.append(g)
    for tc in extras:
        tc.getparent().remove(tc)


def _vmerge_col(tbl, col: int, r0: int, r1: int) -> None:
    """Vertically merge column `col` from row r0..r1 (e.g. logo spanning all header rows).
    Unlike gridSpan, vMerge keeps every cell — the first is 'restart', the rest 'continue'."""
    for r in range(r0, r1 + 1):
        tcPr = tbl.cell(r, col)._tc.get_or_add_tcPr()
        for ex in tcPr.findall(qn("w:vMerge")):
            tcPr.remove(ex)
        vm = OxmlElement("w:vMerge")
        vm.set(qn("w:val"), "restart" if r == r0 else "continue")
        _append_ordered(tcPr, vm, _TCPR_ORDER)


def _resolve_path(rel: str):
    cand = Path(rel)
    if not cand.is_absolute():
        cand = Path(__file__).parent.parent / rel
    return str(cand) if cand.exists() else None


# --- Header styling helpers (close the visual gap vs the hand-made scan) ---

# A4 usable width ~= 16.5cm; original ตัวอย่าง A logo cell ~= 3cm. Keyed by column count.
LOGO_COL_CM = 2.2  # narrow left column holding the vertically-merged logo
# content-column widths (cm) to the RIGHT of the logo column; keyed by content-col count.
HEADER_CONTENT_WIDTHS_CM = {
    1: [14.3],
    2: [9.0, 5.3],
    3: [6.5, 4.3, 3.5],
    4: [5.5, 3.5, 2.8, 2.5],
}
HEADER_CELL_PT = 11  # doc-control text; compact (original scan ~= 11-13pt, body 14pt)

# OOXML child-element sequence (ECMA-376) — insert manually-built elements in schema
# order so Word never triggers a "repair" prompt on open.
_TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders",
    "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription",
]
_TCPR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
    "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
]


def _append_ordered(parent, child, order: list[str]) -> None:
    """Insert `child` into `parent` at its schema-mandated position (per `order`)."""
    pos = order.index(child.tag.split("}")[-1])
    for el in parent:
        local = el.tag.split("}")[-1]
        if local in order and order.index(local) > pos:
            el.addprevious(child)
            return
    parent.append(child)


def _set_table_borders(tbl, size=8, color="000000", val="single") -> None:
    """Visible single-line borders on every cell via w:tblBorders.
    size is eighths of a point: 4 -> 0.5pt, 8 -> 1pt (matches the scanned WI)."""
    tblPr = tbl._tbl.tblPr
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    _append_ordered(tblPr, borders, _TBLPR_ORDER)


def _set_table_widths(tbl, widths_cm: list[float]) -> None:
    """Pin column widths (cm): fixed layout + rewritten tblGrid + per-cell tcW.
    Word lays out fixed tables from tblGrid, so the grid is what truly pins columns."""
    tbl.autofit = False  # writes tblLayout=fixed in correct schema order
    old_grid = tbl._tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl._tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 567)))  # cm -> twips (1cm = 567 twips)
        grid.append(gc)
    tbl._tbl.tblPr.addnext(grid)  # tblGrid sits right after tblPr
    seen = set()
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            tc = cell._tc
            if tc in seen:  # set tcW once per physical cell (skip gridSpan repeats)
                continue
            seen.add(tc)
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _set_cell_margins(cell, top=0.03, bottom=0.03, left=0.1, right=0.1) -> None:
    """Tighten cell text margins (cm) via tcMar — Word's 0.19cm sides look loose vs the scan."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        _append_ordered(tcPr, tcMar, _TCPR_ORDER)
    for tag, val_cm in (("top", top), ("bottom", bottom), ("start", left),
                        ("end", right), ("left", left), ("right", right)):
        el = tcMar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tcMar.append(el)
        el.set(qn("w:w"), str(int(val_cm * 567)))
        el.set(qn("w:type"), "dxa")


def _set_header_cell_format(tbl, has_banner: bool = True) -> None:
    """Vertically center content cells, set 12pt TH Sarabun (+ Thai complex-script
    binding) on every run, tighten margins, compact row heights. Skips banner row 0
    so the org name keeps its larger font and the logo image is left untouched."""
    start = 1 if has_banner else 0  # skip the banner row (logo/org) only when present
    for row in tbl.rows[start:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(HEADER_CELL_PT)
                    run.font.name = THAI_FONT
                    rpr = run._element.get_or_add_rPr()
                    rfonts = rpr.get_or_add_rFonts()
                    rfonts.set(qn("w:cs"), THAI_FONT)  # Thai = complex script
    for i, row in enumerate(tbl.rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.9 if (i == 0 and has_banner) else 0.45)


def _add_field(paragraph, instr: str):
    """Append a Word field (e.g. ' PAGE ', ' NUMPAGES ') that recomputes per page."""
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r.append(begin)
    itxt = OxmlElement("w:instrText")
    itxt.set(qn("xml:space"), "preserve")
    itxt.text = instr
    r.append(itxt)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(end)
    return run


def _shrink_paragraph(p):
    """Collapse an empty spacer paragraph to ~1pt so it doesn't pad the page header."""
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(1)
    return p


def build_header_from_spec(doc: Document, meta: dict[str, str]) -> bool:
    """Data-driven header builder — scales across hospitals (layout from frontmatter).

    Frontmatter:
      header_logo: <path>      (real logo image, embedded)
      header_org:  <str>       (banner org name; falls back to hospital_org)
      header_rows: [<str>...]  each "cell | cell | cell" (pipe-separated).
                               Rows with fewer cells than max → last cell spans.
                               Cells with "label : value" → label rendered bold.

    Returns True if a spec header was built, else False (caller falls back).
    """
    rows = meta.get("header_rows")
    if not rows:
        return False
    if isinstance(rows, str):
        rows = [rows]

    org = meta.get("header_org") or meta.get("hospital_org", "")
    logo_file = _resolve_path(meta["header_logo"]) if meta.get("header_logo") else None

    n_content = max([len(r.split("|")) for r in rows] + [1])
    logo_cols = 1 if logo_file else 0
    total_cols = logo_cols + n_content
    n_rows_total = (1 if org else 0) + len(rows)

    # Build the doc-control box inside the Word PAGE HEADER so it repeats on every page
    # (matches the scanned original). The logo occupies a narrow left column vertically
    # merged down every row; header_rows fill the wider columns to its right (so long
    # cells like doc-type / เรื่อง don't wrap and the box stays compact).
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    tbl = header.add_table(rows=n_rows_total, cols=total_cols, width=Cm(16.5))
    tbl.style = "Table Grid"

    if logo_file:
        lp = tbl.cell(0, 0).paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            lp.add_run().add_picture(logo_file, width=Cm(1.5))
        except Exception:
            lp.add_run("[LOGO]").italic = True

    ridx = 0
    if org:
        op = tbl.cell(0, logo_cols).paragraphs[0]
        op.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = op.add_run(org)
        run.bold = True
        run.font.size = Pt(12)  # org name = largest header text, kept compact
        run.font.name = THAI_FONT
        op.paragraph_format.space_before = Pt(0)
        op.paragraph_format.space_after = Pt(0)
        if total_cols - 1 > logo_cols:
            _merge_range(tbl, 0, logo_cols, total_cols - 1)  # org spans the content columns
        ridx = 1

    for rowstr in rows:
        cells = [c.strip() for c in rowstr.split("|")]
        for j, cell_text in enumerate(cells):
            p = tbl.cell(ridx, logo_cols + j).paragraphs[0]
            if " : " in cell_text or cell_text.split(":")[0].strip() in (
                "ระเบียบปฏิบัติ", "เรื่อง", "เอกสารเลขที่", "รหัสเอกสาร",
                "หน้าที่", "หน้า", "แก้ไขครั้งที่", "ทบทวนแก้ไขครั้งที่",
                "วันที่เริ่มใช้", "วันที่บังคับใช้", "ผู้จัดทำ", "ผู้ตรวจสอบ",
                "ผู้อนุมัติ", "หน่วยงาน",
            ):
                label, _, val = cell_text.partition(":")
                p.add_run(label.strip() + " : ").bold = True
                page_sep = "/" if "/" in val else (" ของ " if " ของ " in val else None)
                if label.strip() in ("หน้าที่", "หน้า") and page_sep:
                    # current page = dynamic PAGE field; total = STATIC from frontmatter
                    # total_pages — faithful to "หน้า N จาก 46" / "หน้า N ของ M". (NUMPAGES
                    # would wrongly show the rendered page count, not the declared total.)
                    total = val.split(page_sep, 1)[1].strip()
                    _add_field(p, " PAGE ")
                    p.add_run(page_sep + total)
                else:
                    p.add_run(val.strip())
            else:
                p.add_run(cell_text)
        if len(cells) < n_content:
            _merge_range(tbl, ridx, logo_cols + len(cells) - 1, total_cols - 1)
        ridx += 1

    if logo_file:
        _vmerge_col(tbl, 0, 0, n_rows_total - 1)  # logo spans every row on the left

    fallback = [(16.5 - LOGO_COL_CM * logo_cols) / n_content] * n_content
    widths = ([LOGO_COL_CM] if logo_cols else []) + HEADER_CONTENT_WIDTHS_CM.get(n_content, fallback)
    _set_table_widths(tbl, widths)
    _set_header_cell_format(tbl, has_banner=bool(org))  # font, valign, margins, heights
    _set_table_borders(tbl)       # visible 1pt black borders on every cell

    # Header housekeeping: shrink the default empty paragraph above the box + add a tiny
    # trailing paragraph (a table must not be the last block — keeps OOXML valid), then set
    # the top margin to hug the compact box (tight gap, no body overlap).
    _shrink_paragraph(header.paragraphs[0])
    _shrink_paragraph(header.add_paragraph())
    section.header_distance = Cm(0.4)
    section.top_margin = Cm(min(max(0.9 + (n_rows_total - 1) * 0.64 + 0.5, 3.0), 7.5))
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    return True


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    set_default_font(doc)

    text = md_path.read_text(encoding="utf-8")
    meta, body = parse_frontmatter(text)
    apply_hospital_profile(meta)  # resolve per-hospital layout profile → header_rows/logo/org

    if meta:
        # Header precedence:
        #   1. header_rows  → data-driven editable header (logo embedded) — scales per hospital
        #   2. header_image → cropped scan embedded as image (pixel-identical, not editable)
        #   3. else         → rebuild generic hospital-b-style tables
        if build_header_from_spec(doc, meta):
            pass
        elif meta.get("header_image") and _resolve_path(meta["header_image"]):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(_resolve_path(meta["header_image"]), width=Cm(16.5))
            doc.add_paragraph("")
        else:
            build_header_tables(doc, meta)

    render_body(doc, body)
    doc.save(docx_path)


def main() -> int:
    if len(sys.argv) != 3:
        print(__doc__)
        return 2
    md = Path(sys.argv[1])
    out = Path(sys.argv[2])
    if not md.exists():
        print(f"ERROR: {md} does not exist")
        return 1
    convert(md, out)
    print(f"OK: {md.name} -> {out.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
