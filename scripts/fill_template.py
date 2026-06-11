#!/usr/bin/env python3
"""
Fill a Word template with markdown body content. [DEPRECATED — see scripts/README.md]

⚠️ SUPERSEDED by `fill_with_docxtpl.py` (mark template once with mark_template.py →
fill with docxtpl placeholders + native numbering + warning box + page footer). This
basic version (doc-control box via paragraph.text, no docxtpl/numbering) is kept for
reference only. For scratch-build with no template, use `md_to_docx.py`.

Preserves template header (logo banner, doc control box, signatures, etc.).
Locates anchor heading "1. วัตถุประสงค์" in template and replaces everything
from that point to end-of-body with content from markdown file.

Usage:
    python scripts/fill_template.py <template.docx> <content.md> <output.docx>
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

# Force UTF-8 stdout so we can print Thai text on Windows cp1252 consoles
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass


ANCHOR_PATTERNS = [
    "1. วัตถุประสงค์",
    "วัตถุประสงค์",
]


def find_anchor(doc):
    for pattern in ANCHOR_PATTERNS:
        for para in doc.paragraphs:
            if pattern in para.text:
                return para
    return None


def delete_from_anchor_to_sectpr(doc, anchor_para) -> int:
    """Delete all body elements from anchor_para onwards, keeping w:sectPr."""
    body = doc.element.body
    all_elements = list(body)
    anchor_el = anchor_para._element

    start_idx = all_elements.index(anchor_el)
    end_idx = len(all_elements)
    for i, el in enumerate(all_elements):
        if el.tag == qn("w:sectPr"):
            end_idx = i
            break

    removed = 0
    for el in all_elements[start_idx:end_idx]:
        body.remove(el)
        removed += 1
    return removed


def _unquote(s: str) -> str:
    s = s.strip()
    if len(s) >= 2 and s[0] == s[-1] and s[0] in ("'", '"'):
        return s[1:-1]
    return s


def parse_frontmatter(text: str) -> tuple[dict[str, str], str]:
    """Parse YAML-style frontmatter. Returns (metadata_dict, body_text)."""
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 4)
    if end == -1:
        return {}, text
    fm_text = text[4:end].strip()
    body = text[end + 4 :].lstrip("\n")

    meta: dict[str, str] = {}
    for line in fm_text.splitlines():
        if ":" in line:
            key, _, value = line.partition(":")
            meta[key.strip()] = _unquote(value)
    return meta, body


def update_doc_control_box(doc, meta: dict[str, str]) -> None:
    """Update section header doc-control table from frontmatter metadata.

    Targets `doc.sections[0].header.tables[0]` (hospital banner + control box).
    Structure observed in ตัวอย่าง B template:
      Row 0: banner (logo + org name)        → KEEP as-is (same hospital)
      Row 1: 3 paragraphs                    → UPDATE
        p[0] = "วิธีปฏิบัติเรื่อง : <subject>"
        p[1] = subject continuation (cleared)
        p[2] = "รหัสเอกสาร : <code>    หน้าที่ : <p>/<total>"
      Row 2: 1 paragraph                     → UPDATE
        p[0] = "ทบทวนแก้ไขครั้งที่ : <N> วันที่ <date>"

    Approach (a): replace paragraph.text — loses run-level formatting but
    inherits paragraph style (Thai font preserved if style defines it).
    Missing meta keys → keep template value (no overwrite).
    """
    if not doc.sections or not doc.sections[0].header.tables:
        return
    tbl = doc.sections[0].header.tables[0]

    # Row 1: subject + doc code + page
    if len(tbl.rows) >= 2:
        cell = tbl.rows[1].cells[0]
        paras = cell.paragraphs

        if len(paras) >= 1 and meta.get("doc_subject"):
            paras[0].text = f"วิธีปฏิบัติเรื่อง : {meta['doc_subject']}"

        # Clear continuation line (p[1]) if subject doesn't overflow
        if len(paras) >= 2:
            paras[1].text = ""

        if len(paras) >= 3 and meta.get("doc_code"):
            page = meta.get("page", "1")
            total = meta.get("total_pages", "[X]")
            # Preserve original whitespace padding pattern for alignment
            paras[2].text = (
                f"รหัสเอกสาร   : {meta['doc_code']}"
                f"                                                            "
                f"หน้าที่ :{page}/{total}"
            )

    # Row 2: revision info
    if len(tbl.rows) >= 3:
        cell = tbl.rows[2].cells[0]
        if cell.paragraphs and meta.get("revision_no") is not None:
            rev_no = meta.get("revision_no", "0")
            rev_date = meta.get("revision_date", "[DD/MM/YY]")
            cell.paragraphs[0].text = (
                f"ทบทวนแก้ไขครั้งที่ : {rev_no} วันที่ {rev_date}"
            )


def is_table_separator(cells: list[str]) -> bool:
    return bool([c for c in cells if c.strip()]) and all(re.match(r"^:?-+:?$", c.strip()) for c in cells if c.strip())


def cells_from(line: str) -> list[str]:
    parts = [p.strip() for p in line.split("|")]
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return parts


def _add_heading_plain(doc, text: str, size_pt: int) -> None:
    """Add a heading as plain paragraph with bold run — does not rely on template styles."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)


def _safe_table_style(table) -> None:
    """Try Table Grid; if template lacks the style, leave default."""
    try:
        table.style = "Table Grid"
    except KeyError:
        pass


def render_markdown(doc, md_text: str) -> None:
    """Render markdown body into doc using plain paragraphs with manual formatting.

    Avoids `add_heading()` and `style="List Bullet"` because templates may lack
    those styles. Instead uses bold + sized runs.
    """
    table_buffer: list[list[str]] = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        cols = max(len(r) for r in table_buffer)
        t = doc.add_table(rows=len(table_buffer), cols=cols)
        _safe_table_style(t)
        for r_idx, row in enumerate(table_buffer):
            for c_idx in range(cols):
                cell_text = row[c_idx] if c_idx < len(row) else ""
                t.cell(r_idx, c_idx).text = cell_text
        table_buffer = []

    for raw in md_text.splitlines():
        line = raw.rstrip("\n")

        if line.lstrip().startswith("|") and "|" in line[1:]:
            cells = cells_from(line)
            if is_table_separator(cells):
                continue
            table_buffer.append(cells)
            continue

        if table_buffer:
            flush_table()

        if line.startswith("#### "):
            _add_heading_plain(doc, line[5:].strip(), size_pt=13)
        elif line.startswith("### "):
            _add_heading_plain(doc, line[4:].strip(), size_pt=14)
        elif line.startswith("## "):
            _add_heading_plain(doc, line[3:].strip(), size_pt=16)
        elif line.startswith("# "):
            _add_heading_plain(doc, line[2:].strip(), size_pt=18)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(f"• {line[2:].strip()}")
        elif line.startswith("---"):
            doc.add_paragraph("")
        else:
            cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            if line.strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(cleaned)

    flush_table()


def main() -> int:
    if len(sys.argv) != 4:
        print(__doc__)
        return 2

    template_path = Path(sys.argv[1])
    content_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])

    if not template_path.exists():
        print(f"ERROR: template not found: {template_path}")
        return 1
    if not content_path.exists():
        print(f"ERROR: content not found: {content_path}")
        return 1

    doc = Document(template_path)

    # Parse frontmatter and use metadata to update section header doc-control box
    raw_text = content_path.read_text(encoding="utf-8")
    meta, md_text = parse_frontmatter(raw_text)
    if meta:
        print(f"Frontmatter: {len(meta)} keys — updating doc control box")
        update_doc_control_box(doc, meta)

    anchor = find_anchor(doc)
    if anchor is None:
        print("WARNING: anchor 'วัตถุประสงค์' not found — appending content at end")
    else:
        print(f"Found anchor: {anchor.text[:60]!r}")
        removed = delete_from_anchor_to_sectpr(doc, anchor)
        print(f"Deleted {removed} elements (body section), preserved header")

    # Render the markdown body back into the doc (after the deleted section,
    # before w:sectPr). Without this the output body is empty.
    render_markdown(doc, md_text)
    print(f"Rendered markdown body ({len(md_text.splitlines())} lines)")

    doc.save(output_path)
    print(f"Saved: {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
