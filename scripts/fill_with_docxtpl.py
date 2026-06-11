#!/usr/bin/env python3
"""
Production-grade WI generator using docxtpl placeholders + 6-proposal synthesis.

Workflow:
  1. Template (.docx) has `{{ placeholders }}` marked by user (one-time, per hospital)
  2. Markdown content file has YAML frontmatter with placeholder values + body content
  3. This script:
     a. docxtpl renders placeholders (anywhere in doc — header, footer, body) from frontmatter
     b. python-docx finds anchor "1. วัตถุประสงค์" in body → deletes body section
     c. Renders markdown body content as new paragraphs (Heading 1/2/3 → Navigation Pane)
     d. Appends "หน้า X จาก Y" page footer + auto-detects AI-disclaimer → wraps as warning box

Body rendering applies:
  - Native Word multilevel numbering (4.1, 4.1.1, 4.1.1.1) via template's numId definitions
  - Word heading styles (Heading 1/2/3) — enables Navigation Pane + TOC outline
  - szCs (Thai complex-script size) paired with sz so headings render correctly in Thai
  - Pagination protections: keep_with_next, cantSplit (rows), widow_control

Usage:
    python scripts/fill_with_docxtpl.py <template_marked.docx> <content.md> <output.docx> [--append]

By default the body anchor ("1. วัตถุประสงค์") MUST be found in the template; if it is
missing the script FAILS HARD (nonzero exit) instead of silently appending new content
after the old template body. Pass --append to opt into the legacy append-at-end behavior.
"""
from __future__ import annotations

import re
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docxtpl import DocxTemplate


if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass


# Allowed placeholder keys — used to filter frontmatter into docxtpl context.
ALLOWED_PLACEHOLDER_KEYS = {
    "doc_subject",
    "doc_code",
    "page",
    "total_pages",
    "rev_no",
    "rev_date",
    "hospital_org",
    "effective_date",
    "author_name",
    "reviewer_name",
    "approver_name",
}

ANCHOR_PATTERNS = ["1. วัตถุประสงค์", "วัตถุประสงค์"]

# Markdown heading level → Word built-in style name. NOTE: # and ## BOTH map to Heading 1
# (## is our top-level section), ### → Heading 2, #### → Heading 3 — so # and ## render at
# the same Navigation-Pane outline level.
_HEADING_STYLE_BY_LEVEL = {1: "Heading 1", 2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}

# Native multilevel numbering — prefers numId definitions in the marked hospital template.
# numId=19 (abstractNum=1) starts at 4 → for procedure clause "4. ขั้นตอนการทำงาน"
# numId=17 (abstractNum=1, override start=1) → fallback for clauses 5+
# These are the IDS the curated hospital templates ship with; in a plain .docx that only
# defines numId 1-9 they do NOT exist, so _resolve_num_id() falls back to an existing numId
# (or None → skip native numbering) instead of emitting a dangling reference.
NUM_ID_BY_CLAUSE = {4: 19}
NUM_ID_DEFAULT = 17
_NUMBERED_PATTERN = re.compile(r"^(\d+)((?:\.\d+)+)\s+(.*)$")

# Indent step (used only for bullet items / leading-whitespace lines without numbering)
INDENT_STEP_PT = 18  # ~6 mm per level


# ----------------------------------------------------------------------------
# Frontmatter + context
# ----------------------------------------------------------------------------

def _unquote(s: str) -> str:
    s = s.strip()
    if len(s) >= 2 and s[0] == s[-1] and s[0] in ("'", '"'):
        return s[1:-1]
    return s


def parse_frontmatter(text: str) -> tuple[dict[str, str], str]:
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 4)
    if end == -1:
        return {}, text
    fm_text = text[4:end].strip()
    body = text[end + 4 :].lstrip("\n")
    meta: dict[str, str] = {}
    for line in fm_text.splitlines():
        if ":" in line:
            key, _, value = line.partition(":")
            meta[key.strip()] = _unquote(value)
    return meta, body


def build_context(meta: dict[str, str]) -> dict[str, str]:
    ctx = {}
    for key in ALLOWED_PLACEHOLDER_KEYS:
        ctx[key] = meta.get(key, "[TBD]")
    if meta.get("revision_no"):
        ctx["rev_no"] = meta["revision_no"]
    if meta.get("revision_date"):
        ctx["rev_date"] = meta["revision_date"]
    return ctx


# ----------------------------------------------------------------------------
# Body section management
# ----------------------------------------------------------------------------

def find_anchor(doc):
    for pattern in ANCHOR_PATTERNS:
        for para in doc.paragraphs:
            if pattern in para.text:
                return para
    return None


def delete_from_anchor_to_sectpr(doc, anchor_para) -> int:
    body = doc.element.body
    all_elements = list(body)
    anchor_el = anchor_para._element
    start_idx = all_elements.index(anchor_el)
    end_idx = len(all_elements)
    for i, el in enumerate(all_elements):
        if el.tag == qn("w:sectPr"):
            end_idx = i
            break
    removed = 0
    for el in all_elements[start_idx:end_idx]:
        body.remove(el)
        removed += 1
    return removed


# ----------------------------------------------------------------------------
# Thai-aware run sizing (proposal 02)
# ----------------------------------------------------------------------------

def _set_run_cs_size(run, size_pt: int) -> None:
    """Pair sz (latin) with szCs (Thai/complex-script) + add cs hint.

    Without szCs, Thai glyphs in a sized run inherit Normal's szCs=32 (16pt),
    so a 14pt heading shows mixed 14pt latin + 16pt Thai — visually broken.
    """
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:hint"), "cs")
    sz_cs = rpr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        rpr.append(sz_cs)
    sz_cs.set(qn("w:val"), str(size_pt * 2))


# ----------------------------------------------------------------------------
# Heading styles (proposal 06)
# ----------------------------------------------------------------------------

def _ensure_heading_style(doc, name: str):
    """Return style by name, instantiating latent built-in if needed. None on failure."""
    try:
        return doc.styles[name]
    except KeyError:
        try:
            return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH, builtin=True)
        except Exception:
            return None


def _add_heading_plain(
    doc,
    text: str,
    size_pt: int,
    space_before_pt: int = 0,
    space_after_pt: int = 0,
    level: int | None = None,
    page_break_before: bool = False,
):
    """Bold heading with optional Word built-in style (gives Navigation Pane / TOC outline).

    Combines:
    - Word heading style application (proposal 06) for outline level
    - Thai-aware szCs sizing (proposal 02) so Thai text matches latin size
    - Pagination protections (keep_with_next, widow_control, optional page_break_before)
    """
    p = doc.add_paragraph()
    if level is not None:
        style = _ensure_heading_style(doc, _HEADING_STYLE_BY_LEVEL.get(level, "Heading 1"))
        if style is not None:
            try:
                p.style = style
            except Exception:
                pass
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    _set_run_cs_size(run, size_pt)
    # Override heading style's default color (often blue) to black — matches WI convention
    run.font.color.rgb = RGBColor(0, 0, 0)
    if space_before_pt:
        p.paragraph_format.space_before = Pt(space_before_pt)
    if space_after_pt:
        p.paragraph_format.space_after = Pt(space_after_pt)
    if page_break_before:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True
    return p


# ----------------------------------------------------------------------------
# Tables
# ----------------------------------------------------------------------------

def _safe_table_style(table) -> None:
    try:
        table.style = "Table Grid"
    except KeyError:
        pass


def _set_table_row_no_split(table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def _apply_widow_control(paragraph) -> None:
    paragraph.paragraph_format.widow_control = True


# ----------------------------------------------------------------------------
# Native multilevel numbering (proposal 01)
# ----------------------------------------------------------------------------

def _existing_num_ids(doc) -> set[int]:
    """Return the set of w:numId values actually defined in the document's numbering part.

    A plain .docx may only define numId 1-9; the curated hospital templates define 17/19.
    Referencing a numId that isn't here yields broken/blank numbering in Word, so callers
    resolve the desired numId against this set first.
    """
    try:
        numbering_el = doc.part.numbering_part.element
    except (NotImplementedError, KeyError, AttributeError):
        return set()
    ids: set[int] = set()
    for num in numbering_el.findall(qn("w:num")):
        val = num.get(qn("w:numId"))
        if val is not None:
            try:
                ids.add(int(val))
            except ValueError:
                pass
    return ids


def _resolve_num_id(doc, clause: int) -> int | None:
    """Pick a numId that EXISTS in `doc` for this clause, or None to skip native numbering.

    Preference order:
      1. The clause's configured numId (e.g. 19) if the template actually defines it.
      2. The default numId (17) if defined.
      3. None → caller keeps the literal "4.1" text via manual numbering.

    Only the template's configured numIds (NUM_ID_BY_CLAUSE / NUM_ID_DEFAULT) are trusted to
    start at the right clause level. We do NOT borrow an arbitrary existing list definition:
    it may be single-level, lack ilvl=1, or start at 1 → blank or "1.x" headings instead of
    "4.x". When neither configured numId is defined, fall back to manual (literal) numbering.
    """
    available = _existing_num_ids(doc)
    if not available:
        return None
    preferred = NUM_ID_BY_CLAUSE.get(clause, NUM_ID_DEFAULT)
    if preferred in available:
        return preferred
    if NUM_ID_DEFAULT in available:
        return NUM_ID_DEFAULT
    return None  # no known-compatible numId → manual numbering preserves the source text


def _apply_native_numbering(paragraph, num_id: int, ilvl: int) -> None:
    """Attach Word multilevel numbering to this paragraph (auto-renumbers on edit).

    `num_id` MUST be a numId that exists in the document (see _resolve_num_id).
    """
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:numPr")):
        pPr.remove(old)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_el = OxmlElement("w:numId")
    num_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(num_el)
    pPr.append(numPr)
    # Numbering owns the indent — clear any manual left_indent
    paragraph.paragraph_format.left_indent = None


def _maybe_render_numbered(doc, line: str):
    """If `line` starts with N.N(.N)* AND clause has a configured numId, render as native list.

    Template's numId=19 starts at 4 (for procedure clause 4). Other clauses (5, 6, 7)
    have no start-aware numId definition, so applying numId=17 (start=1) would render
    "5.1.1" markdown as "1.1.1" in Word. To avoid that bug, skip native numbering for
    unconfigured clauses — they fall through to manual indent.
    """
    m = _NUMBERED_PATTERN.match(line)
    if not m:
        return None
    clause = int(m.group(1))
    if clause not in NUM_ID_BY_CLAUSE:
        return None  # No start-aware numId for this clause; use manual indent instead
    num_id = _resolve_num_id(doc, clause)
    if num_id is None:
        return None  # Target doc defines no usable numId; fall back to manual indent
    depth = m.group(2).count(".")  # "4.1" → 1, "4.1.1" → 2, "4.1.1.1" → 3
    body_text = m.group(3).strip()
    p = doc.add_paragraph(body_text)
    _apply_native_numbering(p, num_id=num_id, ilvl=depth)
    _apply_widow_control(p)
    return p


def _apply_numbered_indent(paragraph, text: str, leading_spaces: int) -> None:
    """Indent paragraph by numbered hierarchy depth (N.N.N) or leading whitespace.

    Priority:
    1. Numbered pattern `^N.N(.N)*` → indent by (dots - 1)
    2. Else leading whitespace (2 spaces per level)

    Used for paragraphs that don't get native numbering (clauses 5+ or non-numbered text).
    """
    num_match = re.match(r"^(\d+\.\d+(?:\.\d+)*)\s+", text)
    if num_match:
        depth = num_match.group(1).count(".") - 1  # "5.1" → 0, "5.1.1" → 1, "5.1.1.1" → 2
        if depth > 0:
            paragraph.paragraph_format.left_indent = Pt(INDENT_STEP_PT * depth)
        return
    if leading_spaces > 0:
        paragraph.paragraph_format.left_indent = Pt(INDENT_STEP_PT * (leading_spaces // 2))


# ----------------------------------------------------------------------------
# Warning box (proposal 05)
# ----------------------------------------------------------------------------

def _set_cell_border(cell, color_hex: str, size: str = "16") -> None:
    """Apply uniform border on all 4 sides via w:tcBorders. size in 1/8 pt."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex.lstrip("#"))
        borders.append(b)
    tc_pr.append(borders)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    tc_pr.append(shd)


def add_warning_box(
    doc,
    text: str,
    border_hex: str = "CC0000",
    fill_hex: str = "FFF0F0",
    label: str = "WARNING",
) -> None:
    """Render `text` inside a 1x1 bordered + shaded table — visual callout."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    _set_cell_border(cell, border_hex, size="16")
    _set_cell_shading(cell, fill_hex)

    cell.text = ""
    p_label = cell.paragraphs[0]
    run_label = p_label.add_run(f"⚠️  {label}: ")
    run_label.bold = True
    run_label.font.size = Pt(12)
    _set_run_cs_size(run_label, 12)
    run_label.font.color.rgb = RGBColor.from_string(border_hex.lstrip("#"))

    body_lines = [ln.rstrip() for ln in text.splitlines() if ln.strip()]
    if body_lines:
        first = body_lines[0]
        r = p_label.add_run(first)
        r.font.size = Pt(11)
        _set_run_cs_size(r, 11)
        for ln in body_lines[1:]:
            p = cell.add_paragraph(ln)
            for r in p.runs:
                r.font.size = Pt(11)
                _set_run_cs_size(r, 11)

    doc.add_paragraph("")  # spacer after box


# ----------------------------------------------------------------------------
# Page footer (proposal 03)
# ----------------------------------------------------------------------------

def add_page_footer(doc, total_pages_label: str = "จาก") -> int:
    """Append right-aligned 'หน้า {PAGE} จาก {NUMPAGES}' paragraph to every section's footer.

    Preserves any existing footer content (e.g. ตัวอย่าง B's signatory block).
    Word recalculates fields on open / on F9.
    """
    touched = 0
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(2)

        def _field(instr: str) -> None:
            r = OxmlElement("w:r")
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            instr_text = OxmlElement("w:instrText")
            instr_text.set(qn("xml:space"), "preserve")
            instr_text.text = instr
            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(qn("w:fldCharType"), "separate")
            t_cache = OxmlElement("w:t")
            t_cache.text = "1"
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            for child in (fld_begin, instr_text, fld_sep, t_cache, fld_end):
                r.append(child)
            p._p.append(r)

        p.add_run("หน้า ")
        _field(" PAGE   \\* MERGEFORMAT ")
        p.add_run(f" {total_pages_label} ")
        _field(" NUMPAGES   \\* MERGEFORMAT ")
        touched += 1
    return touched


# Table of Contents: intentionally NOT generated. The F9-to-populate UX confused MT
# users, so navigation relies on Heading 1/2/3 styles (View → Navigation Pane) instead.
# (Previous insert_toc() helper removed as dead code.)


# ----------------------------------------------------------------------------
# Markdown helpers
# ----------------------------------------------------------------------------

def is_table_separator(cells: list[str]) -> bool:
    return bool([c for c in cells if c.strip()]) and all(re.match(r"^:?-+:?$", c.strip()) for c in cells if c.strip())


def cells_from(line: str) -> list[str]:
    parts = [p.strip() for p in line.split("|")]
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return parts


_WARNING_TRIGGER = re.compile(r"^(⚠|\*\*\s*AI-assisted)", re.IGNORECASE)


def render_markdown(doc, md_text: str) -> None:
    """Render markdown body. Applies headings + native numbering + warning box + pagination."""
    table_buffer: list[list[str]] = []
    warning_buffer: list[str] = []
    in_warning = False
    prev_was_empty = False

    def flush_table() -> None:
        nonlocal table_buffer, prev_was_empty
        if not table_buffer:
            return
        cols = max(len(r) for r in table_buffer)
        t = doc.add_table(rows=len(table_buffer), cols=cols)
        _safe_table_style(t)
        for r_idx, row in enumerate(table_buffer):
            for c_idx in range(cols):
                cell_text = row[c_idx] if c_idx < len(row) else ""
                t.cell(r_idx, c_idx).text = cell_text
        _set_table_row_no_split(t)
        table_buffer = []
        # Spacer after table — keep_with_next so next content stays glued (no orphan heading)
        empty = doc.add_paragraph("")
        empty.paragraph_format.keep_with_next = True
        prev_was_empty = True

    def flush_warning() -> None:
        nonlocal warning_buffer, in_warning, prev_was_empty
        if warning_buffer:
            body = "\n".join(warning_buffer).lstrip()
            # Strip leading ⚠️ + optional bold markers + label prefix (e.g. "AI-assisted draft —")
            body = re.sub(r"^[⚠️\s]*\*+\s*", "", body, count=1)
            body = re.sub(r"\*+\s*$", "", body, flags=re.MULTILINE)
            add_warning_box(doc, body)
            warning_buffer = []
            prev_was_empty = True
        in_warning = False

    def add_empty_if_needed() -> None:
        """Emit a spacer paragraph with keep_with_next so it stays glued to the
        following content (heading or paragraph). Prevents orphan headings —
        e.g. '7. การควบคุมคุณภาพ' alone at page bottom with 7.1 on next page.
        """
        nonlocal prev_was_empty
        if not prev_was_empty:
            empty = doc.add_paragraph("")
            empty.paragraph_format.keep_with_next = True
            prev_was_empty = True

    for raw in md_text.splitlines():
        line = raw.rstrip("\n")
        stripped = line.lstrip()
        leading_spaces = len(line) - len(stripped)

        # Warning box detection
        if not in_warning and _WARNING_TRIGGER.match(stripped):
            if table_buffer:
                flush_table()
            in_warning = True
            warning_buffer.append(stripped)
            continue

        if in_warning:
            # End on next `#` heading or `---` rule
            if stripped.startswith("#") or stripped.startswith("---"):
                flush_warning()
                # fall-through: process current line as normal below
            else:
                warning_buffer.append(stripped)
                continue

        # Table row
        if stripped.startswith("|") and "|" in stripped[1:]:
            cells = cells_from(stripped)
            if is_table_separator(cells):
                continue
            table_buffer.append(cells)
            prev_was_empty = False
            continue

        if table_buffer:
            flush_table()

        # Headings
        if stripped.startswith("#### "):
            _add_heading_plain(doc, stripped[5:].strip(), size_pt=13,
                               space_before_pt=8, space_after_pt=4, level=4)
            prev_was_empty = False
        elif stripped.startswith("### "):
            _add_heading_plain(doc, stripped[4:].strip(), size_pt=14,
                               space_before_pt=12, space_after_pt=6, level=3)
            prev_was_empty = False
        elif stripped.startswith("## "):
            _add_heading_plain(doc, stripped[3:].strip(), size_pt=16,
                               space_before_pt=16, space_after_pt=8, level=2,
                               page_break_before=True)
            prev_was_empty = False
        elif stripped.startswith("# "):
            _add_heading_plain(doc, stripped[2:].strip(), size_pt=18,
                               space_before_pt=20, space_after_pt=10, level=1,
                               page_break_before=True)
            prev_was_empty = False
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(f"• {stripped[2:].strip()}")
            _apply_numbered_indent(p, stripped[2:].strip(), leading_spaces)
            _apply_widow_control(p)
            prev_was_empty = False
        elif stripped.startswith("---"):
            add_empty_if_needed()
        else:
            cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            if not cleaned:
                add_empty_if_needed()
            else:
                # Try native multilevel numbering first (4.1, 4.1.1, 4.1.1.1)
                p = _maybe_render_numbered(doc, cleaned)
                if p is None:
                    p = doc.add_paragraph(cleaned)
                    _apply_numbered_indent(p, cleaned, leading_spaces)
                    _apply_widow_control(p)
                prev_was_empty = False

    flush_table()
    flush_warning()


def main() -> int:
    args = sys.argv[1:]
    allow_append = False
    if "--append" in args:
        allow_append = True
        args = [a for a in args if a != "--append"]
    if len(args) != 3:
        print(__doc__)
        return 2

    template_path = Path(args[0])
    content_path = Path(args[1])
    output_path = Path(args[2])

    if not template_path.exists():
        print(f"ERROR: template not found: {template_path}")
        return 1
    if not content_path.exists():
        print(f"ERROR: content not found: {content_path}")
        return 1

    # Pass 1: docxtpl renders all {{placeholders}} from frontmatter
    raw_text = content_path.read_text(encoding="utf-8")
    meta, md_body = parse_frontmatter(raw_text)
    ctx = build_context(meta)
    print(f"Frontmatter: {len(meta)} keys parsed, {len(ctx)} placeholder values built")

    tpl = DocxTemplate(template_path)
    tpl.render(ctx)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        tpl.save(tmp_path)
        print(f"Pass 1 done: placeholders rendered → {tmp_path.name}")

        # Pass 2: python-docx replaces body section starting at anchor
        doc = Document(tmp_path)
        anchor = find_anchor(doc)
        if anchor is None:
            if not allow_append:
                print(
                    "ERROR: body anchor 'วัตถุประสงค์' not found in template "
                    f"({template_path}). The generated WI would otherwise keep the old "
                    "template body AND append the new content. Mark the anchor in the "
                    "template, or pass --append to force append-at-end behavior."
                )
                return 1
            print("WARNING: anchor 'วัตถุประสงค์' not found — --append set, appending body at end")
        else:
            removed = delete_from_anchor_to_sectpr(doc, anchor)
            print(f"Pass 2: deleted {removed} body elements, will append markdown content")

        # TOC removed — Word's F9-to-populate UX is confusing for MT users.
        # Navigation Pane (View → Navigation Pane) provides outline navigation via
        # Heading 1 styles already applied to `##` sections.
        render_markdown(doc, md_body)
        sections_with_footer = add_page_footer(doc)
        print(f"Pass 3: page footer added to {sections_with_footer} section(s)")

        doc.save(output_path)
    finally:
        tmp_path.unlink(missing_ok=True)  # never orphan the temp file, even on failure

    print(f"Saved: {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
