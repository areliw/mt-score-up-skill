#!/usr/bin/env python3
"""
Personnel replacement utility — swap names across a filled WI docx.

Handles the personnel-rotation scenario (per rule-mt-in-drivers-seat): when a
person leaves/changes role, their name in the template's signature block (footer,
header, or body) must update across all WIs. Template hardcodes names (not
placeholders), so we post-process the rendered docx.

Searches body paragraphs + tables + each section's primary header/footer
(and any explicitly-defined first/even header/footer) paragraphs + tables.

Usage:
    python scripts/replace_personnel.py <doc.docx> "<old name>=<new name>" ["<old2>=<new2>" ...] [--dry-run]

  --dry-run : report replacement count without writing. Otherwise the input docx is
              overwritten in place AND a <doc>.docx.bak backup is written first.

Example:
    python scripts/replace_personnel.py wi.docx "[ชื่อเดิม นามสกุล]=[ชื่อใหม่ นามสกุล]"
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document


if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> int:
    """If paragraph text contains any old string, rebuild paragraph with replacement.

    Collapses runs (loses intra-paragraph formatting) — acceptable for name lines
    where formatting is uniform. Returns count of replacements made.
    """
    full = paragraph.text
    changed = 0
    new_full = full
    for old, new in replacements.items():
        if old in new_full:
            new_full = new_full.replace(old, new)
            changed += 1
    if changed and new_full != full:
        for r in paragraph.runs:
            r.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_full
        else:
            paragraph.add_run(new_full)
    return changed


def replace_in_tables(tables, replacements: dict[str, str]) -> int:
    count = 0
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    count += replace_in_paragraph(p, replacements)
    return count


def replace_all(doc, replacements: dict[str, str]) -> int:
    count = 0
    # Body paragraphs + tables
    for p in doc.paragraphs:
        count += replace_in_paragraph(p, replacements)
    count += replace_in_tables(doc.tables, replacements)
    # Section headers + footers (paragraphs + tables)
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            for p in hf.paragraphs:
                count += replace_in_paragraph(p, replacements)
            count += replace_in_tables(hf.tables, replacements)
    return count


def main() -> int:
    dry_run = "--dry-run" in sys.argv
    args = [a for a in sys.argv[1:] if a != "--dry-run"]
    if len(args) < 2:
        print(__doc__)
        return 2
    doc_path = Path(args[0])
    replacements: dict[str, str] = {}
    for pair in args[1:]:
        if "=" in pair:
            old, _, new = pair.partition("=")
            replacements[old.strip()] = new.strip()
    if not doc_path.exists():
        print(f"ERROR: {doc_path} not found")
        return 1
    if not replacements:
        print("ERROR: no valid 'old=new' pair given (need '=' in each)")
        return 2

    doc = Document(doc_path)
    n = replace_all(doc, replacements)
    for old, new in replacements.items():
        print(f"  '{old}' -> '{new}'")

    if dry_run:
        print(f"DRY RUN — {n} occurrence(s) would change in {doc_path.name} (not saved)")
        return 0

    # This script overwrites its input in place → back up the original first.
    backup = doc_path.with_suffix(doc_path.suffix + ".bak")
    if not backup.exists():
        shutil.copy2(doc_path, backup)
        print(f"Backup saved: {backup.name}")
    doc.save(doc_path)
    print(f"Replaced {n} occurrence(s) in {doc_path.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
